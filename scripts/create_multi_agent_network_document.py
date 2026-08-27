"""Create a focused project document for the implemented multi-agent network."""

from datetime import UTC, datetime
from pathlib import Path

from create_project_document import create_agent_mcp_image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "SAP_SuccessFactors_Multi_Agent_Network_Design.docx"
DIAGRAM = ROOT / "agent_mcp_architecture.png"
NAVY = "17365D"
BLUE = "2563EB"
TEAL = "0F766E"
PALE = "EAF2FF"
GRAY = "F1F5F9"
WHITE = "FFFFFF"


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    properties.append(fill)


def margins(cell, value: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:tcMar")
    for side in ("top", "start", "bottom", "end"):
        item = OxmlElement(f"w:{side}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        node.append(item)
    properties.append(node)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    item = doc.add_table(rows=1, cols=len(headers))
    item.style = "Table Grid"
    item.alignment = WD_TABLE_ALIGNMENT.CENTER
    item.autofit = False
    repeat_header(item.rows[0])
    for index, value in enumerate(headers):
        cell = item.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for row_index, values in enumerate(rows):
        cells = item.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)
            if row_index % 2:
                shade(cell, GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def callout(doc: Document, title: str, body: str) -> None:
    item = doc.add_table(rows=1, cols=1)
    repeat_header(item.rows[0])
    item.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = item.cell(0, 0)
    shade(cell, PALE)
    margins(cell, 170)
    run = cell.paragraphs[0].add_run(title + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    cell.paragraphs[0].add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def bullet(doc: Document, value: str) -> None:
    doc.add_paragraph(value, style="List Bullet")


def build() -> None:
    create_agent_mcp_image()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 28, NAVY), ("Subtitle", 14, TEAL),
                              ("Heading 1", 19, NAVY), ("Heading 2", 13, BLUE),
                              ("Heading 3", 11, TEAL)):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "SAP SUCCESSFACTORS MULTI-AGENT NETWORK  |  TECHNICAL DESIGN"
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Read-only synthetic-data prototype")

    p = doc.add_paragraph("SAP SuccessFactors", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Multi-Agent Network", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Architecture, governance, MCP tools and test-tenant roadmap", style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    callout(doc, "Project status", "Implemented locally with the OpenAI Agents SDK Runner in demo mode, eleven read-only MCP tools, synthetic SuccessFactors-style data, server-side authorization and safe lifecycle tracing.")
    table(doc, ["Item", "Value"], [
        ["Document type", "Technical architecture and project design"],
        ["Version", "1.0"],
        ["Date", datetime.now(UTC).date().isoformat()],
        ["Primary audience", "HRIS, SuccessFactors, security, architecture and engineering teams"],
        ["Current test baseline", "34 passing automated tests; one opt-in paid live smoke test skipped by default"],
    ], [1.8, 5.05])
    doc.add_page_break()

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph("The multi-agent network provides a controlled conversational layer over read-only Employee Central-style data. An HR Orchestrator owns the final response and delegates bounded tasks to Employee Central, Position Management, Data Quality and Security Review specialists. Specialists do not access files, databases or SAP directly; permitted MCP tools are their only data boundary.")
    doc.add_paragraph("Security is deterministic below the model. Application-controlled identity, role authorization, population scope, field allow-lists, result limits, sanitization and audit logging remain authoritative even if a model selects an inappropriate action.")
    callout(doc, "Design principle", "Agents decide how to coordinate approved capabilities. They do not decide who the user is, what data the user may access, or which fields are safe to expose.")

    doc.add_heading("2. Network architecture", level=1)
    doc.add_picture(str(DIAGRAM), width=Inches(6.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    props = doc.inline_shapes[-1]._inline.docPr
    props.set("title", "Implemented multi-agent and MCP architecture")
    props.set("descr", "HR Orchestrator delegates to four specialist agents. Specialists receive bounded MCP tools that enter a shared authorization, sanitization and audit boundary.")
    caption = doc.add_paragraph("Figure 1. Implemented orchestrator, specialist and MCP permission network")
    caption.style = doc.styles["Caption"]
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table(doc, ["Layer", "Component", "Responsibility"], [
        ["Experience", "Streamlit", "Collects identity, effective date and question; displays answer, trace, tools and usage."],
        ["Application", "FastAPI /api/chat", "Creates session/correlation context, applies request policy and selects runtime mode."],
        ["Orchestration", "Agents SDK Runner", "Runs the manager pattern with turn and output-token limits."],
        ["Coordination", "HR Orchestrator", "Delegates bounded subtasks, combines evidence and owns the final response."],
        ["Specialization", "Four specialist agents", "Perform domain-specific retrieval, validation and review."],
        ["Tool boundary", "FastMCP", "Publishes eleven typed, read-only tools and applies per-agent filtering."],
        ["Control", "HRToolService", "Authorizes, limits, sanitizes and audits every provider call."],
        ["Data", "Typed provider", "Uses synthetic data now; preserves a seam for SuccessFactors OData V2."],
    ], [1.1, 1.7, 4.05])

    doc.add_heading("3. Agent catalogue", level=1)
    table(doc, ["Agent", "Primary responsibility", "Data boundary"], [
        ["HR Orchestrator", "Understand intent, select specialists, combine evidence and produce the answer.", "Specialists-as-tools only; no unrestricted HR provider access."],
        ["Employee Central", "Employee, employment, job, organization, manager and effective-date interpretation.", "Five employee/job MCP tools."],
        ["Position Management", "Positions, incumbents, vacancies, hierarchy, foundation objects and metadata.", "Five position/foundation MCP tools."],
        ["Data Quality", "Execute deterministic individual or population validations and explain findings.", "Three validation/history MCP tools."],
        ["Security Review", "Review requests and proposed responses for restricted-data exposure.", "No direct HR MCP access."],
    ], [1.45, 3.0, 2.4])

    doc.add_heading("4. MCP permission matrix", level=1)
    table(doc, ["MCP operation", "EC", "Position", "DQ", "Security"], [
        ["get_employee_basic_profile", "Yes", "-", "-", "-"],
        ["get_employee_job_information", "Yes", "-", "-", "-"],
        ["get_employment_information", "Yes", "-", "-", "-"],
        ["get_manager_relationship", "Yes", "-", "-", "-"],
        ["get_position / search_positions", "-", "Yes", "-", "-"],
        ["get_foundation_object / get_entity_metadata", "-", "Yes", "-", "-"],
        ["validate_employee_data / population", "-", "-", "Yes", "-"],
        ["get_effective_dated_changes", "Yes", "Yes", "Yes", "-"],
    ], [3.25, 0.75, 0.95, 0.75, 1.15])
    doc.add_paragraph("A dynamic MCP filter enforces this matrix independently of prompts. The Security Review specialist intentionally has no HR retrieval access.")

    doc.add_heading("5. Delegation and response sequence", level=1)
    for step in (
        "FastAPI resolves the synthetic user and creates session and correlation identifiers.",
        "The Runner starts the HR Orchestrator with application-controlled context.",
        "The Orchestrator selects one or more specialists based on the question.",
        "Each specialist sees only its permitted MCP tools.",
        "MCP derives identity from protected headers and invokes HRToolService.",
        "HRToolService authorizes population scope before provider access, sanitizes records and writes audit metadata.",
        "Specialists return bounded evidence; the Orchestrator produces the final answer.",
        "Safe hooks publish agent/tool lifecycle names, token usage and estimated cost.",
    ):
        bullet(doc, step)
    callout(doc, "Example compound request", "For E1004, show the current job and assigned position, validate the assignment, and review the result for restricted-data exposure. This can involve Employee Central, Position Management, Data Quality and Security Review before final synthesis.")

    doc.add_heading("6. Identity, security and privacy", level=1)
    table(doc, ["Control", "Current implementation"], [
        ["Application context", "user_id, role, session_id and correlation_id are created outside model input."],
        ["MCP transport", "Identity and an internal MCP credential are passed in HTTP headers, not tool arguments."],
        ["Authorization", "Server-owned tool and employee-population allow-lists are checked before provider access."],
        ["Field security", "Entity allow-lists and explicit denied fields remove unknown or prohibited properties."],
        ["National ID", "Fifty synthetic PerNationalId records exist for testing; nationalId remains denied and no retrieval tool is exposed."],
        ["Injection resistance", "Identity changes, secrets, prohibited data, raw URLs/OData and writes are rejected."],
        ["Trace safety", "Only lifecycle names are captured; prompts, arguments, outputs and private reasoning are excluded."],
        ["Audit", "Correlation, role, tool, outcome, count, duration and source are recorded without full HR payloads."],
    ], [1.55, 5.3])
    callout(doc, "Production warning", "The local user selector is a demonstration identity, not authentication. Production requires enterprise SSO, signed sessions, managed secrets, gateway controls and formal SuccessFactors RBP reconciliation.")

    doc.add_heading("7. Runtime modes and cost controls", level=1)
    table(doc, ["Mode", "Orchestration", "Provider", "Credentials"], [
        ["mock", "Deterministic router", "Synthetic MockProvider", "None"],
        ["demo", "OpenAI Agents SDK Runner", "Synthetic MockProvider through MCP", "OpenAI key + local MCP token"],
        ["real", "Currently deterministic", "Configured ODataProvider", "Tenant OAuth configuration"],
    ], [0.85, 2.1, 2.2, 1.7])
    doc.add_paragraph("Demo cost controls include DEMO_MAX_TURNS, DEMO_MAX_OUTPUT_TOKENS and DEMO_DAILY_TOKEN_BUDGET. Configurable per-million-token rates provide a local estimate. OpenAI tracing is disabled by default and sensitive trace content remains disabled when tracing is enabled.")

    doc.add_heading("8. Test strategy", level=1)
    table(doc, ["Test level", "Purpose"], [
        ["Unit", "Schemas, allow-lists, redaction, OData escaping, data generation and cost calculation."],
        ["Service", "Authorization, population scope, sanitization, audit and deterministic quality rules."],
        ["MCP", "Tool inventory, absence of model-callable identity fields and specialist tool filters."],
        ["API", "Health, answers, effective dates, denials, evidence and compound flows."],
        ["Live smoke", "One opt-in paid request validates Runner delegation, MCP use and usage reporting."],
    ], [1.35, 5.5])
    doc.add_paragraph("Current verified baseline: 34 tests passed and one paid live test was skipped by default; Ruff and Mypy passed.")

    doc.add_heading("9. SuccessFactors test-system roadmap", level=1)
    for step in (
        "Confirm the tenant OData API base URL, company ID, token endpoint and approved test population.",
        "Create a dedicated read-only API user and minimum Role-Based Permissions.",
        "Register an OAuth client and protect its private key outside source control.",
        "Implement ConfiguredOAuthTokenProvider with signed assertion, token caching, expiry and redacted errors.",
        "Validate $metadata, entity names, field types, effective dating, custom fields and Position Management availability.",
        "Map tenant responses into canonical records before sanitization.",
        "Add bounded pagination and aggregate tools; never send entire populations to an agent.",
        "Test direct provider access, authorization denials, timeouts and token expiry before enabling agents.",
        "Separate orchestration selection from provider selection so Agents SDK + SuccessFactors can be enabled explicitly.",
        "Complete privacy, security, audit, retention and production-readiness reviews.",
    ):
        bullet(doc, step)

    doc.add_heading("10. Key implementation files", level=1)
    table(doc, ["File", "Responsibility"], [
        ["src/sf_mcp_poc/agents.py", "Agent definitions, specialist tools, dynamic MCP filter and safe hooks."],
        ["src/sf_mcp_poc/demo_runtime.py", "Runner execution, MCP connection, usage accounting and cost controls."],
        ["src/sf_mcp_poc/mcp_server.py", "Eleven typed read-only MCP tools and protected header context."],
        ["src/sf_mcp_poc/service.py", "Authorization-aware execution, validation, sanitization and audit."],
        ["src/sf_mcp_poc/provider.py", "Synthetic data and future SuccessFactors OData provider seam."],
        ["src/sf_mcp_poc/security.py", "Identities, roles, scopes, allow-lists, redaction and sanitization."],
        ["tests/", "Deterministic, security, agent-runtime and optional live tests."],
    ], [2.65, 4.2])

    doc.core_properties.title = "SAP SuccessFactors Multi-Agent Network Technical Design"
    doc.core_properties.subject = "Orchestrator, specialist agents, MCP tools and security controls"
    doc.core_properties.author = "Project Team"
    doc.core_properties.keywords = "SAP SuccessFactors, multi-agent, MCP, OpenAI Agents SDK, FastAPI"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
