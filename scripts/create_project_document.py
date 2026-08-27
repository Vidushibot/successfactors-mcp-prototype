"""Create the formal project report DOCX."""

from datetime import UTC, datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "SAP_SuccessFactors_Multi-Agent_MCP_Prototype_Project_Document.docx"
ARCHITECTURE_IMAGE = ROOT / "project_architecture.png"
AGENT_MCP_IMAGE = ROOT / "agent_mcp_architecture.png"
NAVY = "17365D"
BLUE = "2563EB"
TEAL = "0F766E"
PALE_BLUE = "EAF2FF"
PALE_TEAL = "E6F6F3"
PALE_GRAY = "F1F5F9"
WHITE = "FFFFFF"


def create_architecture_image() -> None:
    image = Image.new("RGB", (1800, 1040), "white")
    draw = ImageDraw.Draw(image)
    font_dir = Path("C:/Windows/Fonts")
    regular = ImageFont.truetype(str(font_dir / "arial.ttf"), 27)
    small = ImageFont.truetype(str(font_dir / "arial.ttf"), 22)
    bold = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 28)
    title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 38)

    def box(x1, y1, x2, y2, fill, outline, title, subtitle=""):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=outline, width=4)
        center = (x1 + x2) // 2
        title_box = draw.textbbox((0, 0), title, font=bold)
        draw.text(
            (center - (title_box[2] - title_box[0]) // 2, y1 + 20), title, font=bold, fill=outline
        )
        if subtitle:
            lines = subtitle.split("\n")
            for index, line in enumerate(lines):
                bounds = draw.textbbox((0, 0), line, font=small)
                draw.text(
                    (center - (bounds[2] - bounds[0]) // 2, y1 + 62 + index * 28),
                    line,
                    font=small,
                    fill="#334155",
                )

    def arrow(start, end, color="#475569", width=5):
        draw.line((start, end), fill=color, width=width)
        x2, y2 = end
        x1, y1 = start
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 > x1 else -1
            points = [(x2, y2), (x2 - direction * 20, y2 - 12), (x2 - direction * 20, y2 + 12)]
        else:
            direction = 1 if y2 > y1 else -1
            points = [(x2, y2), (x2 - 12, y2 - direction * 20), (x2 + 12, y2 - direction * 20)]
        draw.polygon(points, fill=color)

    title = "Current secure demo-mode architecture"
    bounds = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((1800 - bounds[2]) // 2, 25), title, font=title_font, fill="#17365D")
    draw.text(
        (520, 78),
        "Live Agents SDK orchestration with application-controlled MCP identity",
        font=regular,
        fill="#64748B",
    )

    box(55, 160, 310, 285, "#DBEAFE", "#1D4ED8", "HR user", "Synthetic identity")
    box(385, 160, 640, 285, "#EDE9FE", "#7C3AED", "Streamlit UI", "Chat • evidence • audit")
    box(715, 160, 970, 285, "#E0F2FE", "#0369A1", "FastAPI", "/api/chat • /docs")
    box(
        1045,
        160,
        1390,
        285,
        "#FEF3C7",
        "#B45309",
        "Agents SDK Runner",
        "HR Orchestrator + specialists",
    )
    arrow((310, 222), (385, 222))
    arrow((640, 222), (715, 222))
    arrow((970, 222), (1045, 222))

    box(
        610,
        390,
        1080,
        545,
        "#CCFBF1",
        "#0F766E",
        "HRToolService",
        "Authorization • validation • limits\nSanitization • audit logging",
    )
    arrow((1215, 285), (970, 390))
    box(85, 410, 480, 525, "#FFEDD5", "#C2410C", "FastMCP server", "11 narrow read-only tools")
    arrow((480, 467), (610, 467))

    box(
        630,
        650,
        1060,
        795,
        "#F1F5F9",
        "#334155",
        "Typed provider interface",
        "Same contract for mock and real modes",
    )
    arrow((845, 545), (845, 650))
    box(
        250,
        885,
        710,
        1005,
        "#DCFCE7",
        "#15803D",
        "MockProvider",
        "Synthetic Employee Central-style data",
    )
    box(
        990,
        885,
        1550,
        1005,
        "#F8FAFC",
        "#64748B",
        "ODataProvider",
        "Future SAP test tenant • fail closed",
    )
    arrow((740, 795), (500, 885), "#15803D")
    arrow((950, 795), (1220, 885), "#64748B")

    image.save(ARCHITECTURE_IMAGE)


def create_agent_mcp_image() -> None:
    image = Image.new("RGB", (1800, 1420), "white")
    draw = ImageDraw.Draw(image)
    font_dir = Path("C:/Windows/Fonts")
    small = ImageFont.truetype(str(font_dir / "arial.ttf"), 20)
    body = ImageFont.truetype(str(font_dir / "arial.ttf"), 23)
    bold = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 26)
    title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 38)

    def centered_text(x1, y1, x2, text, font, fill, top):
        bounds = draw.textbbox((0, 0), text, font=font)
        draw.text(((x1 + x2 - (bounds[2] - bounds[0])) // 2, top), text, font=font, fill=fill)

    def box(x1, y1, x2, y2, fill, outline, title, lines):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=outline, width=4)
        centered_text(x1, y1, x2, title, bold, outline, y1 + 18)
        for index, line in enumerate(lines):
            centered_text(x1 + 10, y1, x2 - 10, line, small, "#334155", y1 + 62 + index * 29)

    def arrow(start, end, color="#475569", width=5):
        draw.line((start, end), fill=color, width=width)
        x2, y2 = end
        x1, y1 = start
        if abs(x2 - x1) >= abs(y2 - y1):
            direction = 1 if x2 > x1 else -1
            points = [(x2, y2), (x2 - direction * 20, y2 - 12), (x2 - direction * 20, y2 + 12)]
        else:
            direction = 1 if y2 > y1 else -1
            points = [(x2, y2), (x2 - 12, y2 - direction * 20), (x2 + 12, y2 - direction * 20)]
        draw.polygon(points, fill=color)

    centered_text(
        0,
        0,
        1800,
        "Implemented agent delegation and MCP tool architecture",
        title_font,
        "#17365D",
        28,
    )
    centered_text(
        0,
        0,
        1800,
        "Specialists are bounded tools; the HR Orchestrator owns the final answer",
        body,
        "#64748B",
        82,
    )
    box(
        590,
        140,
        1210,
        270,
        "#F3E8FF",
        "#7C3AED",
        "HR Orchestrator Agent",
        ["Understands the HR question", "Selects specialists • combines evidence • final response"],
    )

    agents = [
        (
            45,
            "Employee Central Agent",
            "#DBEAFE",
            "#1D4ED8",
            [
                "Employee • employment • job • manager",
                "Effective-dated interpretation",
                "Approved MCP tools only",
            ],
        ),
        (
            485,
            "Position Management Agent",
            "#E0F2FE",
            "#0369A1",
            [
                "Positions • incumbents • vacancies",
                "Hierarchy • foundation objects",
                "Bounded position searches",
            ],
        ),
        (
            925,
            "Data Quality Agent",
            "#CCFBF1",
            "#0F766E",
            [
                "Deterministic validation findings",
                "Errors versus warnings",
                "Individual and population checks",
            ],
        ),
        (
            1365,
            "Security Review Agent",
            "#FEE2E2",
            "#B91C1C",
            [
                "Authorization and field review",
                "Injection and secret protection",
                "No independent HR-data access",
            ],
        ),
    ]
    for x, title, fill, outline, lines in agents:
        box(x, 370, x + 390, 540, fill, outline, title, lines)
        arrow((900, 270), (x + 195, 370), outline, 4)

    tool_boxes = [
        (
            45,
            "Employee MCP tools",
            "#EFF6FF",
            "#1D4ED8",
            [
                "get_employee_basic_profile",
                "get_employee_job_information",
                "get_employment_information",
                "get_manager_relationship",
                "get_effective_dated_changes",
            ],
        ),
        (
            485,
            "Position MCP tools",
            "#F0F9FF",
            "#0369A1",
            [
                "get_position",
                "search_positions",
                "get_foundation_object",
                "get_entity_metadata",
                "get_effective_dated_changes",
            ],
        ),
        (
            925,
            "Data-quality MCP tools",
            "#F0FDFA",
            "#0F766E",
            [
                "validate_employee_data",
                "validate_employee_population",
                "get_effective_dated_changes",
                "Uses deterministic rule evidence",
                "No write or correction tools",
            ],
        ),
        (
            1365,
            "Security controls",
            "#FEF2F2",
            "#B91C1C",
            [
                "Reviews planned/completed calls",
                "Checks sanitized evidence",
                "Rejects secrets/raw OData/writes",
                "Respects opaque denials",
                "No unrestricted MCP access",
            ],
        ),
    ]
    for index, (x, title, fill, outline, lines) in enumerate(tool_boxes):
        box(x, 650, x + 390, 875, fill, outline, title, lines)
        arrow((x + 195, 540), (x + 195, 650), agents[index][3], 4)

    box(
        350,
        1010,
        1450,
        1155,
        "#FFEDD5",
        "#C2410C",
        "FastMCP server: narrow read-only SuccessFactors tools",
        [
            "Strict typed schemas • no raw URLs • no arbitrary OData • no writes",
            "Every call enters the same deterministic security boundary",
        ],
    )
    for x in (240, 680, 1120, 1560):
        arrow((x, 875), (900, 1010), "#C2410C", 4)

    box(
        420,
        1260,
        1380,
        1385,
        "#E2E8F0",
        "#334155",
        "HRToolService security boundary",
        ["Identity → authorization → limits → provider → sanitization → audit"],
    )
    arrow((900, 1155), (900, 1260), "#334155", 5)
    image.save(AGENT_MCP_IMAGE)


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    properties.append(fill)


def margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    tc_mar = properties.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        properties.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2:
                shade(cell, PALE_GRAY)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.size = Pt(9)
            if widths:
                cell.width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc: Document, title: str, text: str, color=PALE_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, color)
    margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    run = p.add_run(title + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run = p.add_run(text)
    run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def bullet(doc: Document, text: str, level=0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    paragraph.add_run(text)


def build() -> None:
    create_architecture_image()
    create_agent_mcp_image()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, color in (
        ("Title", 30, NAVY),
        ("Subtitle", 15, TEAL),
        ("Heading 1", 20, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11, TEAL),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name != "Normal" else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.text = "SAP SUCCESSFACTORS MULTI-AGENT MCP PROTOTYPE  |  PROJECT REPORT"
    header.style = styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    doc.add_paragraph("SAP SuccessFactors", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph("Multi-Agent MCP Prototype", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Secure, read-only Employee Central analysis using synthetic SuccessFactors-style data",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    add_callout(
        doc,
        "Prototype notice",
        "Demonstration using synthetic SuccessFactors-style data. This application is not an SAP SuccessFactors system and is not production-ready.",
        PALE_TEAL,
    )
    metadata = doc.add_table(rows=5, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.style = "Table Grid"
    for row, values in enumerate(
        (
            ("Document", "Project report"),
            ("Version", "2.0"),
            ("Date", datetime.now(UTC).date().isoformat()),
            ("Primary mode", "Live Agents SDK demo over synthetic HR data"),
            ("Technology", "Python 3.11, FastAPI, Streamlit, FastMCP, OpenAI Agents SDK"),
        )
    ):
        for col, value in enumerate(values):
            cell = metadata.cell(row, col)
            margins(cell)
            if col == 0:
                shade(cell, NAVY)
                run = cell.paragraphs[0].add_run(value)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                cell.text = value
    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Prepared for local demonstration, coursework, architecture review and controlled testing"
    )
    run.italic = True
    run.font.color.rgb = RGBColor.from_string("64748B")
    doc.add_page_break()

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "The SAP SuccessFactors Multi-Agent MCP Prototype demonstrates how authorized HR users can ask read-only questions about Employee Central-style data while preserving a strict security boundary. The application uses coherent synthetic data by default, requires no SAP credentials, and does not claim to be an SAP system."
    )
    doc.add_paragraph(
        "All HR data access is mediated by narrow typed operations, server-side authorization, field allow-lists, deterministic validation, sanitization and structured audit logging. The provider interface allows the synthetic provider to be replaced later by a tenant-specific SuccessFactors OData V2 provider without redesigning the UI or tool schemas."
    )
    add_callout(
        doc,
        "Current implementation status",
        "When APP_MODE=demo, /api/chat invokes the OpenAI Agents SDK Runner. The HR Orchestrator delegates to bounded specialists that can see only permitted MCP tools. Mock mode remains available for credential-free deterministic testing.",
    )

    doc.add_heading("2. Project objectives and boundaries", level=1)
    add_table(
        doc,
        ["Area", "Implemented approach"],
        [
            [
                "Goal",
                "Evidence-based HR, position and data-quality answers for authorized prototype users.",
            ],
            [
                "Primary users",
                "HR administrators, HRIS analysts, EC consultants, position specialists and data-quality teams.",
            ],
            [
                "Data source",
                "Synthetic SuccessFactors-style Employee Central records in mock mode.",
            ],
            [
                "Operations",
                "Read-only retrieval, search, effective dating, validation and bounded population analysis.",
            ],
            [
                "Explicit exclusions",
                "Writes, arbitrary HTTP, raw OData clauses, Basic Authentication and prohibited HR fields.",
            ],
            [
                "Production status",
                "Educational proof of concept; demonstration identity and SQLite are not production controls.",
            ],
        ],
        [1.35, 5.5],
    )

    doc.add_heading("3. Architecture", level=1)
    doc.add_picture(str(ARCHITECTURE_IMAGE), width=Inches(6.8))
    picture_properties = doc.inline_shapes[-1]._inline.docPr
    picture_properties.set(
        "descr",
        "Current demo architecture: HR user to Streamlit, FastAPI, Agents SDK Runner, secured HR tool service, typed provider interface, mock provider and future OData provider.",
    )
    picture_properties.set("title", "Current secure demo-mode architecture")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph(
        "Figure 1. Current secure demo architecture and provider replacement seam"
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.style = styles["Caption"]
    doc.add_heading("3.1 Current demo-mode runtime", level=2)
    add_table(
        doc,
        ["Step", "Component", "Responsibility"],
        [
            [
                "1",
                "Authorized HR user",
                "Selects a synthetic identity and submits a read-only HR question.",
            ],
            [
                "2",
                "Streamlit UI",
                "Collects the question/effective date and displays answers, evidence, tools, trace and audit events.",
            ],
            [
                "3",
                "FastAPI backend",
                "Validates request schemas and exposes health, configuration, chat and audit APIs.",
            ],
            [
                "4",
                "Agents SDK Runner",
                "Runs the HR Orchestrator with maximum-turn and output-token controls.",
            ],
            [
                "5",
                "Specialist agents",
                "Delegate Employee Central, Position, Data Quality and Security Review work.",
            ],
            [
                "6",
                "FastMCP + HRToolService",
                "Enforces tool filtering, identity, authorization, limits, sanitization and audit.",
            ],
            [
                "7",
                "Typed provider",
                "Uses synthetic data now and preserves a seam for a future ODataProvider.",
            ],
        ],
        [0.55, 1.65, 4.65],
    )
    doc.add_heading("3.2 MCP boundary", level=2)
    doc.add_paragraph(
        "The FastMCP server exposes eleven narrow read-only tools. Identity, session and correlation values arrive in application-controlled HTTP headers and are absent from model-callable schemas. No tool accepts raw URLs, arbitrary entity names or model-generated OData fragments."
    )
    doc.add_heading("3.3 Implemented live-agent path", level=2)
    doc.add_paragraph(
        "The HR Orchestrator invokes four bounded specialists as tools and retains final-response ownership. A dynamic MCP filter enforces least privilege per specialist. Authenticated identity and authorization context are controlled by the application and are never sourced from prompt text."
    )
    doc.add_picture(str(AGENT_MCP_IMAGE), width=Inches(6.8))
    agent_picture_properties = doc.inline_shapes[-1]._inline.docPr
    agent_picture_properties.set(
        "descr",
        "Implemented manager-agent architecture: HR Orchestrator delegates to Employee Central, Position Management, Data Quality and Security Review specialists. Each specialist uses only relevant narrow FastMCP tools before the shared HRToolService security boundary.",
    )
    agent_picture_properties.set("title", "Implemented agent delegation and MCP tool architecture")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    agent_caption = doc.add_paragraph(
        "Figure 2. Implemented specialist delegation and relevant MCP tool assignments"
    )
    agent_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    agent_caption.style = styles["Caption"]
    doc.add_heading("3.4 Technology stack", level=2)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose in the prototype"],
        [
            [
                "Language/runtime",
                "Python 3.11+",
                "Typed application, services, providers, validation and automation.",
            ],
            [
                "User interface",
                "Streamlit 1.62",
                "Synthetic user selection, chat, effective dates, evidence, traces and audit viewing.",
            ],
            [
                "Application API",
                "FastAPI + Uvicorn",
                "Typed HTTP endpoints, OpenAPI documentation, health and chat APIs.",
            ],
            [
                "Agent framework",
                "OpenAI Agents SDK",
                "Runner execution, orchestrator ownership, specialists-as-tools, hooks and usage metrics.",
            ],
            [
                "MCP",
                "Official MCP Python SDK + FastMCP",
                "Eleven narrow read-only HR tools using strict server-side controls.",
            ],
            [
                "HTTP/OData",
                "httpx",
                "Controlled future calls to SuccessFactors OData V2 with explicit timeouts.",
            ],
            [
                "Schemas/settings",
                "Pydantic v2 + pydantic-settings",
                "Strict inputs/outputs, environment validation and fail-closed real mode.",
            ],
            [
                "Persistence",
                "SQLite + SQLAlchemy 2.x",
                "Prototype user rules and structured sanitized audit events.",
            ],
            [
                "Testing",
                "pytest + pytest-asyncio",
                "Credential-free unit, integration, security and regression testing.",
            ],
            ["Quality", "Ruff + Mypy", "Formatting, linting and useful static type checking."],
            [
                "Local configuration",
                "python-dotenv",
                "Development-only loading of placeholder environment configuration.",
            ],
            [
                "Packaging/deployment",
                "pyproject.toml, PowerShell, Docker",
                "Windows-first native execution and optional development containers.",
            ],
        ],
        [1.25, 1.8, 3.8],
    )
    add_callout(
        doc,
        "Credential model",
        "Mock mode requires no SAP or OpenAI credentials. Demo mode requires OPENAI_API_KEY and a local MCP_INTERNAL_TOKEN. Real mode requires validated non-production SAP OAuth configuration and never falls back silently to mock data.",
        PALE_TEAL,
    )

    doc.add_heading("4. Specialist agents", level=1)
    add_table(
        doc,
        ["Specialist", "Responsibilities", "Representative operations", "Must never do"],
        [
            [
                "Employee Central",
                "Employee, employment, job, manager, organization and effective-date interpretation.",
                "Basic profile; job and employment information; manager relationship; effective-dated changes.",
                "Retrieve prohibited fields, exceed population scope, invent facts or perform writes.",
            ],
            [
                "Position Management",
                "Position details, incumbents, vacancies, hierarchy and organizational placement.",
                "Get/search positions; foundation objects; position counts and missing-assignment analysis.",
                "Treat missing incumbent as proven vacancy without evidence or accept raw OData.",
            ],
            [
                "Data Quality",
                "Explain deterministic rule findings and separate errors from warnings.",
                "Validate one employee or an authorized employee population.",
                "Invent client-specific rules, correct data or upgrade warnings without evidence.",
            ],
            [
                "Security Review",
                "Review requests and evidence for authorization, injection and prohibited-data exposure.",
                "Approved-field verification; safe rejection; permission and write checks.",
                "Retrieve HR data merely to grant access, expand permissions or reveal hidden prompts/secrets.",
            ],
        ],
        [1.25, 2.05, 2.0, 1.55],
    )
    doc.add_heading("4.1 Live delegation example", level=2)
    doc.add_paragraph(
        "For a compound E1004 review in demo mode, the Runner can execute this bounded delegation sequence. The visible trace contains lifecycle metadata only, never chain-of-thought:"
    )
    for text in (
        "HR Orchestrator Agent starts",
        "Employee Central specialist -> get_employee_job_information",
        "Position Management specialist -> get_position",
        "Data Quality specialist -> validate_employee_data",
        "Security Review specialist -> approved-field verification",
        "HR Orchestrator -> final evidence-based answer",
    ):
        bullet(doc, text)

    doc.add_heading("5. Security and authorization model", level=1)
    doc.add_paragraph(
        "Security is enforced deterministically below the conversational layer. Agent instructions reinforce policy but are not the authorization mechanism."
    )
    add_table(
        doc,
        ["Control", "Implementation"],
        [
            [
                "Identity",
                "Four synthetic users resolved server-side; explicitly unsuitable for production authentication.",
            ],
            [
                "Tool authorization",
                "Each role has a server-owned allow-list of permitted operations.",
            ],
            [
                "Population scope",
                "Restricted users can access only explicitly assigned synthetic employees.",
            ],
            [
                "Field controls",
                "Central entity registry and field allow-lists; unknown and denied fields are dropped.",
            ],
            [
                "OData safety",
                "Typed values, escaped literals, selected fields, capped $top and rejection of URLs/query fragments.",
            ],
            [
                "Existence protection",
                "Unauthorized and unavailable records use an opaque public denial.",
            ],
            [
                "Audit",
                "Correlation, user, role, tool, outcome, record count, duration and source are stored without full payloads.",
            ],
            [
                "Prompt injection",
                "Requests for credentials, hidden prompts, raw OData, writes or prohibited data are rejected.",
            ],
        ],
        [1.55, 5.3],
    )
    add_callout(
        doc,
        "Prohibited data",
        "Compensation, bank information, national/government identifiers, medical and disability information, emergency contacts, dependents, home address, personal email/phone, date of birth, gender and marital status are explicitly denied.",
        PALE_TEAL,
    )

    doc.add_heading("6. Synthetic data and supported analysis", level=1)
    add_table(
        doc,
        ["Dataset", "Volume", "Coverage"],
        [
            ["Employees", "10", "E1001-E1010 with fictional names and employment records."],
            ["Job records", "14", "Historical, current, overlapping and future-dated scenarios."],
            ["Positions", "12", "P100-P111 with incumbents, vacancies and parent relationships."],
            ["Companies", "2", "Synthetic foundation objects."],
            ["Business units", "3", "Synthetic foundation objects."],
            ["Divisions", "4", "Synthetic foundation objects."],
            ["Departments", "6", "Synthetic foundation objects."],
        ],
        [1.7, 0.8, 4.35],
    )
    doc.add_heading("6.1 Intentional quality scenarios", level=2)
    add_table(
        doc,
        ["Identifier", "Scenario"],
        [
            ["E1001", "Valid current employee with historical job change."],
            ["E1002", "Overlapping effective-dated job records."],
            ["E1003", "Future-dated record that overlaps the current open-ended record."],
            ["E1004", "Assigned position P999 does not exist."],
            ["E1006", "Manager E9998 does not exist."],
            ["E1007", "Department mismatch and invalid foundation-object assignment."],
            ["E1008", "Terminated employee scenario."],
            ["E1010", "Missing position."],
            ["P109", "Only position without a parent position."],
        ],
        [1.2, 5.65],
    )
    doc.add_paragraph(
        "Exact exported records are available in test_data/employees.json, jobs.json, positions.json and foundation_objects.json. TEST_DATA.md documents the scenarios."
    )

    doc.add_heading("7. Interfaces and user experience", level=1)
    doc.add_heading("7.1 FastAPI endpoints", level=2)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /", "Application notice and navigation."],
            ["GET /health and /ready", "Liveness and readiness."],
            ["GET /api/config/public", "Safe public mode and read-only status."],
            ["GET /api/demo-users", "Synthetic roles and authorized scopes."],
            [
                "POST /api/chat",
                "Question, effective date, evidence, tools, trace and correlation ID.",
            ],
            ["GET /api/audit/events", "Role-filtered audit-event access."],
        ],
        [2.05, 4.8],
    )
    doc.add_heading("7.2 Streamlit experience", level=2)
    for item in (
        "Synthetic user selector and role/scope summary",
        "Effective-date selection and example questions",
        "Direct evidence-based answers",
        "Tools-used, safe delegation-trace, token-usage and estimated-cost panels",
        "Sanitized evidence records and references",
        "Authorization-denial messaging and correlation IDs",
        "Role-filtered audit-event viewer",
    ):
        bullet(doc, item)

    doc.add_heading("8. Testing and verification", level=1)
    doc.add_paragraph(
        "At the time this report was generated, the latest verification reported 33 passing automated tests and one intentionally skipped paid live test, plus successful Ruff and Mypy checks."
    )
    add_table(
        doc,
        ["Test area", "Representative coverage"],
        [
            [
                "Core",
                "Entity registry, allow-lists, sanitization, redaction, OData escaping and record limits.",
            ],
            ["Authorization", "Role tools, restricted population and non-disclosure behavior."],
            [
                "API",
                "Health, direct answers, effective dates, audit persistence and compound flows.",
            ],
            [
                "Data quality",
                "Missing/invalid position, overlaps, manager absence and population scans.",
            ],
            [
                "Security",
                "Prompt injection, prohibited fields, raw OData and absence of write endpoints.",
            ],
            ["MCP", "Tool inventory, hidden identity fields and specialist-specific allow-lists."],
            [
                "Live agents",
                "Opt-in Runner smoke test verifies delegation, MCP use and token accounting.",
            ],
        ],
        [1.45, 5.4],
    )
    add_callout(
        doc,
        "Testing principle",
        "The normal suite never calls OpenAI. Set RUN_LIVE_AGENT_TESTS=true only when intentionally running one paid live delegation smoke test against running local services.",
    )

    doc.add_heading("9. Running the prototype on Windows", level=1)
    doc.add_paragraph(
        "Copy .env.example to .env. For demo mode set APP_MODE=demo, OPENAI_MODEL, OPENAI_API_KEY and a random MCP_INTERNAL_TOKEN. Then run:"
    )
    p = doc.add_paragraph()
    p.style = styles["No Spacing"]
    run = p.add_run(
        "Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass\n.\\scripts\\setup_windows.ps1\n.\\scripts\\run_ui_windows.ps1"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    add_table(
        doc,
        ["Service", "URL"],
        [
            ["Streamlit UI", "http://127.0.0.1:8501"],
            ["FastAPI", "http://127.0.0.1:8000"],
            ["API documentation", "http://127.0.0.1:8000/docs"],
            ["MCP", "http://127.0.0.1:8001/mcp"],
        ],
        [2.0, 4.85],
    )
    doc.add_paragraph(
        "The Windows UI launcher starts MCP and FastAPI automatically when unavailable. Keep the Streamlit terminal open. Never commit the .env file or expose API keys in prompts, traces or screenshots."
    )

    doc.add_heading("10. Limitations and roadmap", level=1)
    add_table(
        doc,
        ["Current limitation", "Recommended next step"],
        [
            [
                "Live model output is probabilistic and can vary between runs.",
                "Add evaluation datasets and thresholds for routing, answer accuracy and grounded citations.",
            ],
            [
                "The local daily token ledger resets on process restart.",
                "Use a durable organization-level budget and usage alerting for shared deployments.",
            ],
            [
                "Real SAP OAuth is intentionally not configured.",
                "Validate the exact tenant flow against SAP documentation and a non-production tenant.",
            ],
            [
                "Synthetic identity is locally selectable.",
                "Replace with SSO, signed sessions and enterprise authorization mapping.",
            ],
            [
                "SQLite is a prototype audit store.",
                "Use a managed, access-controlled and monitored audit platform.",
            ],
            [
                "Position and workforce analytics are bounded prototypes.",
                "Add approved aggregate tools with scope-aware counts and explicit definitions.",
            ],
            [
                "Docker Compose is development-only.",
                "Add hardened images, secrets management, network policy and production observability.",
            ],
        ],
        [2.85, 4.0],
    )

    doc.add_heading("11. Definition of success", level=1)
    for item in (
        "The application runs locally in mock mode without SAP or OpenAI credentials.",
        "Only narrow read-only tools are available.",
        "Server-side authorization and population restrictions are enforced before provider access.",
        "Prohibited and unknown fields are removed before evidence reaches the conversational layer.",
        "Answers identify the data source and effective date and include evidence references.",
        "Every tool invocation creates a sanitized audit event.",
        "Unauthorized access does not disclose whether a record exists.",
        "Automated checks pass without weakening security controls.",
        "Demo requests report safe delegation events and token usage without exposing private reasoning.",
    ):
        bullet(doc, item)

    doc.add_heading("Appendix A. Key project files", level=1)
    add_table(
        doc,
        ["File", "Purpose"],
        [
            ["README.md", "Setup, modes, commands, troubleshooting and limitations."],
            ["ARCHITECTURE.md", "Detailed component, sequence and trust-boundary design."],
            ["SECURITY.md", "Threat model, controls, risks and production backlog."],
            ["TEST_DATA.md", "Synthetic records and intentional quality scenarios."],
            [
                "SUCCESSFACTORS_ARCHITECTURE.excalidraw",
                "Editable current/target architecture diagram.",
            ],
            ["src/sf_mcp_poc/api.py", "FastAPI routes and mock/demo execution selection."],
            [
                "src/sf_mcp_poc/agents.py",
                "Orchestrator, specialists, tool filters and safe trace hooks.",
            ],
            [
                "src/sf_mcp_poc/demo_runtime.py",
                "Agents SDK Runner, MCP context, usage and cost controls.",
            ],
            ["src/sf_mcp_poc/service.py", "Authorization-aware tool execution and validation."],
            ["src/sf_mcp_poc/mcp_server.py", "Typed read-only FastMCP tool surface."],
            ["src/sf_mcp_poc/provider.py", "Mock/OData provider abstraction and query builder."],
            ["tests/", "Unit, integration and security regression tests."],
        ],
        [2.75, 4.1],
    )

    doc.core_properties.title = "SAP SuccessFactors Multi-Agent MCP Prototype Project Report"
    doc.core_properties.subject = "Secure read-only Employee Central MCP prototype"
    doc.core_properties.author = "Project Team"
    doc.core_properties.keywords = (
        "SAP SuccessFactors, Employee Central, MCP, FastAPI, Streamlit, agents"
    )
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
